from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).with_name('App2_Assignment_and_Research_Workspaces_Functional_Specification.docx')
NAVY = '202A33'; GOLD = 'C89B2C'; PALE_GOLD = 'F6EFD9'; LIGHT = 'F2F4F7'; MID = '5D6670'; GREEN = '26734D'; RED = 'A53636'; WHITE = 'FFFFFF'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

def font(run, size=11, bold=False, color='222222', italic=False):
    run.font.name = 'Calibri'; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = RGBColor.from_string(color)
    return run

styles = doc.styles
normal = styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string('222222')
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,before,after in [('Heading 1',16,16,8),('Heading 2',13,12,6),('Heading 3',12,8,4)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(GOLD if name!='Heading 3' else NAVY); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for name in ['List Bullet','List Number']:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(11); s.paragraph_format.left_indent=Inches(.5); s.paragraph_format.first_line_indent=Inches(-.25); s.paragraph_format.space_after=Pt(5); s.paragraph_format.line_spacing=1.1

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.LEFT
font(header.add_run('PUBLIC SERVICE COMMISSION  |  APP2 FUNCTIONAL SPECIFICATION'),9,True,MID)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run('Assignment and Research Workspaces  |  Controlled Working Draft'),9,False,MID)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def borders(table, color='D8DCE1'):
    tblPr=table._tbl.tblPr; el=tblPr.find(qn('w:tblBorders'))
    if el is None: el=OxmlElement('w:tblBorders'); tblPr.append(el)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        x=OxmlElement(f'w:{edge}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'4'); x.set(qn('w:color'),color); el.append(x)

def set_cell_width(cell, dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(dxa)); tcW.set(qn('w:type'),'dxa')

def mark_header(row):
    trPr=row._tr.get_or_add_trPr(); flag=OxmlElement('w:tblHeader'); flag.set(qn('w:val'),'true'); trPr.append(flag)

def table(headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; borders(t); mark_header(t.rows[0])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; set_cell_width(c,widths[i]); shade(c,NAVY); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(h),9,True,WHITE)
    for r,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            set_cell_width(cells[i],widths[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r%2: shade(cells[i],'FAFAFA')
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(2); font(p.add_run(str(val)),9.3,False,'222222')
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def h(text,level=1): doc.add_heading(text,level=level)
def p(text,bold_lead=None):
    x=doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        font(x.add_run(bold_lead),11,True,NAVY); font(x.add_run(text[len(bold_lead):]),11)
    else: font(x.add_run(text),11)
    return x
def bullets(items):
    for item in items: font(doc.add_paragraph(item,style='List Bullet').runs[0],11)
def numbers(items):
    for item in items: font(doc.add_paragraph(item,style='List Number').runs[0],11)
def callout(label,text,color=PALE_GOLD):
    t=doc.add_table(rows=1,cols=1); t.autofit=False; mark_header(t.rows[0]); set_cell_width(t.cell(0,0),9360); shade(t.cell(0,0),color); borders(t,GOLD)
    q=t.cell(0,0).paragraphs[0]; q.paragraph_format.space_after=Pt(2); font(q.add_run(label+'  '),10,True,GOLD); font(q.add_run(text),10.5,False,NAVY)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

# Cover
doc.add_paragraph().paragraph_format.space_after=Pt(70)
k=doc.add_paragraph(); k.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(k.add_run('APP2 PRODUCT AND OPERATING MODEL'),10,True,GOLD)
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; t.paragraph_format.space_after=Pt(8); font(t.add_run('Assignment and Research Workspaces'),28,True,NAVY)
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER; s.paragraph_format.space_after=Pt(60); font(s.add_run('Functional specification, workflow definition and implementation guide'),14,False,MID)
meta=doc.add_table(rows=4,cols=2); meta.alignment=WD_TABLE_ALIGNMENT.CENTER; meta.autofit=False; borders(meta); mark_header(meta.rows[0])
for i,(a,b) in enumerate([('System','App2 - Assignment and Knowledge Management System'),('Document purpose','Define the complete operational model for Assignment and Research Workspaces'),('Audience','Product owners, managers, researchers, reviewers, developers and testers'),('Status','Controlled working specification - August 2026')]):
    set_cell_width(meta.cell(i,0),2200); set_cell_width(meta.cell(i,1),7160); shade(meta.cell(i,0),LIGHT); font(meta.cell(i,0).paragraphs[0].add_run(a),10,True,NAVY); font(meta.cell(i,1).paragraphs[0].add_run(b),10)
doc.add_page_break()

h('1. Executive Summary')
p('App2 requires two connected but distinct operational environments. The Assignment Workspace controls delivery of defined work: responsibility, tasks, deadlines, outputs, review and completion. The Research Workspace controls disciplined inquiry: research questions, methodology, evidence, analysis, research outputs and approval.')
callout('CORE DESIGN RULE','Assignments are delivery-led; research projects are evidence-led. They may be linked, but neither workspace should become a renamed version of the other.')
h('1.1 Intended outcome',2)
bullets(['Users can understand the current state of work immediately.','Every action has an accountable owner, deadline and audit trail.','Only the selected workspace tab is rendered and actionable.','Required evidence and approvals cannot be bypassed.','Notifications lead to a valid outstanding action and clear automatically when the action is completed.','Felix assists users without silently modifying official records.'])
h('1.2 Shared workspace experience',2)
p('Both workspaces should open full screen and use the same App2 command-centre pattern: compact identity header, Back and Close controls, persistent tabs, KPI summary, next-action guidance, role-aware actions, restrained charcoal-and-gold styling, responsive layouts and controlled evidence.')

h('2. Conceptual Difference')
table(['Dimension','Assignment Workspace','Research Workspace'],[
('Primary purpose','Deliver a defined instruction or output','Answer a defined question using defensible evidence'),
('Starting point','Assignment brief','Research problem and question'),
('Core unit of work','Task','Milestone, source, analysis and report section'),
('Planning emphasis','Who will do what and by when','Why, how, with what evidence and under what method'),
('Primary controls','Task ownership, deadlines, status and review','Methodology, evidence provenance, analysis, citations and research review'),
('Completion test','Requested deliverables approved and assignment closed','Research plan, evidence, analysis and final report approved'),
('Typical output','Brief, memo, report, action or operational deliverable','Research plan, instruments, dataset, findings, brief and final research report')],[1700,3600,4060])

h('3. Shared Functional Principles')
numbers(['Identity first: show a human-readable reference, title, division or research domain, status, lead and timeline.','Immediate navigation: selecting a tab must instantly activate it and render only that panel.','State-derived guidance: Next Action must come from actual workflow state.','Role-aware controls: hide or disable actions the current user cannot perform.','No false evidence: never display unrelated records merely to fill a panel.','Controlled completion: mandatory tasks, milestones, documents and reviews must be satisfied before closure.','Auditability: record material changes with actor, timestamp, action and affected record.','Responsive operation: full functionality must remain usable on desktop, tablet and mobile.'])

h('4. Assignment Workspace Functional Specification')
h('4.1 Purpose',2)
p('The Assignment Workspace is the operational command centre for a single assignment. It coordinates delivery from setup through planning, execution, submission, review, corrections, approval and completion.')
h('4.2 Header and navigation',2)
bullets(['Assignment title and human-readable reference.','Division, status, priority, due date and days remaining.','Back and Close controls.','Tabs: Overview, Tasks, Team, Documents & Outputs, Discussion, Activity and Review.','Bright but restrained status and deadline indicators.'])
h('4.3 Overview tab',2)
bullets(['Workflow stage and explanatory wizard text.','KPI strip for status, priority, completion, due date, lead and team size.','Assignment brief and expected outcome.','Current tasks and progress visible without excessive scrolling.','Next Action card with a real action button.','Team preview and recent activity.'])
h('4.4 Tasks tab',2)
bullets(['Create tasks with title, description, owner, priority, mandatory due date, start date and notes.','Default task deadline from the assignment deadline while allowing an earlier date.','Filter by All, My Tasks, Not Started, In Progress, Completed and Overdue.','Display owner, deadline health, progress and status without overlap.','Deadline colors: green when safe, amber when approaching and red when urgent, overdue or missing.','Only managers may allocate work; task owners may update permitted status and progress.','Completion updates assignment progress and resolves applicable notifications.'])
h('4.5 Team tab',2)
bullets(['Identify assignment lead, contributors and reviewer responsibilities.','Add or remove members subject to permissions.','Use consistent avatars, aligned names and role labels.','Show task counts, workload and completion by member.','Prevent unassigned responsibility where governance requires an owner.'])
h('4.6 Documents & Outputs tab',2)
bullets(['Upload working files and download attachments.','Create controlled documents from approved templates.','Assign document sections to team members.','Track versions, section completion, review comments and final outputs.','Separate working material from formally approved deliverables.','Never show unrelated evidence or generic documents as assignment support.'])
h('4.7 Discussion, Activity and Review tabs',2)
bullets(['Discussion supports updates, questions, mentions and recorded decisions.','Activity provides a chronological read-only audit of assignment events.','Review supports submission, reviewer assignment, start review, request changes, resubmission and approval.','Review comments are required for requested changes.','Approved assignments display a clear completed-review state.'])

h('5. Assignment Workflow and Controls')
table(['Stage','Primary actor','Required action','Exit condition'],[
('1. Set Up','Manager','Create brief, division, deadline, priority and initial team','Valid assignment record exists'),
('2. Plan','Manager / Lead','Create tasks, owners and due dates','Required work is allocated'),
('3. Work','Task owners','Perform tasks and create working outputs','Mandatory tasks and outputs are ready'),
('4. Submit','Lead','Submit completed assignment for formal review','Submission event recorded'),
('5. Review','Reviewer','Assess evidence and deliverables','Decision recorded'),
('6. Corrections','Lead / Team','Complete requested changes and resubmit','Corrections addressed'),
('7. Approved','Reviewer / Manager','Approve deliverables','Approval recorded'),
('8. Complete','Manager','Verify closure conditions and close assignment','Assignment locked as complete')],[1300,1700,3560,2800])
callout('COMPLETION GATE','An assignment cannot be completed while mandatory tasks are incomplete, required outputs are missing, or the latest review is not approved.')

h('6. Research Workspace Functional Specification')
h('6.1 Purpose',2)
p('The Research Workspace is the evidence and methodology command centre for a research project. It should manage the full lifecycle from question definition and plan approval through evidence collection, analysis, drafting, review and approved publication.')
h('6.2 Header and navigation',2)
bullets(['Research title, reference, status, lead researcher and timeline.','Linked assignment where applicable.','Back and Close controls.','Tabs: Overview, Research Plan, Team, Sources, Documents, Discussion, Report and Activity.','Counters for sources, documents and unresolved collaboration activity.'])
h('6.3 Overview tab',2)
bullets(['Project summary and decision context.','KPI strip for milestone progress, sources, documents and approved report sections.','Research Readiness checklist.','Deadline and risk summary.','Dynamic Next Action directing the user to the missing research requirement.','Recent evidence, team and activity preview.'])
h('6.4 Research Plan tab',2)
bullets(['Problem statement and primary research question.','Objectives, scope, exclusions and expected outputs.','Methodology, target population, sampling and data-collection methods.','Ethical, privacy, confidentiality and information-security considerations.','Assumptions, constraints and quality criteria.','Dated milestones with owners and statuses.','Formal plan approval before controlled evidence collection begins.'])
h('6.5 Team tab',2)
bullets(['Assign a lead researcher, contributors, analyst and reviewer.','Assign ownership of milestones, evidence collections and report sections.','Show workload, overdue responsibilities and completed contributions.','Restrict sensitive datasets and evidence to authorized roles.'])
h('6.6 Sources and evidence tab',2)
bullets(['Capture journals, reports, policy documents, legislation, datasets, books, websites, interviews and field evidence.','Record author, institution, publisher, publication date, identifier, URL and notes.','Classify source provenance as approved App2 evidence, external evidence or primary field evidence.','Assess source quality, relevance and recency.','Detect duplicates and incomplete metadata.','Link sources to findings and report sections.','Generate consistent citations and a reference list.'])
h('6.7 Documents and report tabs',2)
bullets(['Create research plans, instruments, analysis notes, briefs and reports from governed templates.','Assign document and report sections to individual team members.','Track section status: Not Started, Draft, In Progress, Ready for Review and Approved.','Maintain controlled versions and review records.','Link supporting sources and citations to sections.','Compile approved sections into a final report.','Prevent final approval when required sections or citations are incomplete.'])
h('6.8 Discussion and Activity tabs',2)
bullets(['Record questions, updates, methodological decisions and evidence discussions.','Mention team members and resolve threads.','Maintain chronological activity for plan changes, evidence additions, section edits and approvals.','Keep the audit history read-only to normal users.'])

h('7. Research Workflow and Controls')
table(['Stage','Primary actor','Required action','Exit condition'],[
('1. Initiation','Manager / Lead','Define problem, question, sponsor and linked assignment','Research project accepted'),
('2. Plan','Lead / Team','Define scope, objectives, methodology, ethics and milestones','Research plan complete'),
('3. Plan Approval','Manager / Reviewer','Assess feasibility, method and safeguards','Plan approved'),
('4. Evidence Collection','Researchers','Capture controlled sources, instruments and field evidence','Evidence threshold met'),
('5. Analysis','Lead / Analyst','Code, compare, validate and document findings','Findings supported'),
('6. Drafting','Section owners','Draft report with linked evidence and citations','Required sections ready'),
('7. Review','Reviewer','Assess methodology, evidence, findings and conclusions','Decision recorded'),
('8. Corrections','Research team','Address review issues and resubmit','Corrections verified'),
('9. Approved / Complete','Manager / Reviewer','Approve final report and close project','Final outputs controlled')],[1350,1700,3560,2750])
callout('RESEARCH COMPLETION GATE','Research cannot be completed until the plan is approved, required evidence is recorded, mandatory report sections are approved and the final controlled output exists.')

h('8. Roles and Permissions')
table(['Role','Assignment responsibilities','Research responsibilities'],[
('Administrator','System configuration, exceptional access and audit oversight','System configuration, governed access and audit oversight'),
('Research Manager','Create assignments, allocate work, monitor, review and close','Create projects, appoint lead, approve plan, monitor and close'),
('Research Officer','Complete owned tasks, contribute documents and submit work','Collect evidence, complete milestones, analyse and draft sections'),
('Reviewer','Review submitted deliverables and record decisions','Review plan, methodology, evidence and final report'),
('Lead researcher','Where assigned: coordinate assignment delivery','Coordinate plan, team, evidence, analysis, report and submission')],[1600,3850,3910])

h('9. Notifications and Next Actions')
p('Notifications must represent current outstanding work. Opening a notification should navigate directly to the relevant workspace and tab. Once the user performs the required action, the notification must cease to appear as urgent and must not continue pointing to an already completed action.')
table(['Trigger','Recipient','Target','Resolution'],[
('Task assigned','Task owner','Assignment / Tasks','Task acknowledged or completed according to policy'),
('Task overdue','Owner and lead','Assignment / Tasks','Due date corrected or task completed'),
('Assignment submitted','Reviewer','Assignment / Review','Review started or decision recorded'),
('Research plan submitted','Manager / Reviewer','Research / Research Plan','Plan decision recorded'),
('Milestone approaching','Owner and lead','Research / Research Plan','Milestone completed or rescheduled'),
('Report section ready','Reviewer','Research / Report','Section decision recorded'),
('Changes requested','Lead and responsible owners','Relevant Review or Report panel','Corrected resubmission recorded')],[2200,2000,2200,2960])

h('10. Felix Assistance and Safeguards')
bullets(['Recognize the current App2 system and workspace context.','Explain assignment or research status using the correct record.','Suggest next actions based on actual workflow state.','Summarize only authorized evidence.','Identify evidence gaps, unsupported claims and missing citations.','Assist drafting through suggestions that require explicit user acceptance.','Never change official records, approve work or assign responsibility without an authorized user action.','When approved evidence is unavailable, state the limitation instead of substituting unrelated records.'])

h('11. Data and Audit Requirements')
bullets(['Stable unique IDs plus human-readable references.','Created, updated, submitted, reviewed, approved and completed timestamps.','Actor identity for every material state transition.','Role and permission checks enforced by the backend, not only the interface.','Version histories for controlled documents and reports.','Source provenance, identifiers and evidence links.','Immutable activity records for governance events.','Retention and archival rules for final outputs and sensitive research data.'])

h('12. Usability and Visual Requirements')
bullets(['Full-screen workspace with no hidden sidebar dependency.','Compact professional density; avoid oversized empty cards.','Only one selected tab panel visible or rendered at a time.','No duplicated legacy workspace markup.','Status colors remain readable and accessible: green for healthy/complete, amber for approaching risk and red for urgent/blocked/overdue.','Tables and progress indicators must not overlap at supported widths.','Forms clearly identify mandatory fields, validation errors and save state.','Mobile layouts preserve actions and avoid horizontal content loss.','Unsupported fields are hidden from ordinary users rather than exposing developer-facing limitations.'])

h('13. Acceptance Criteria')
numbers(['Opening either workspace covers the full application viewport and retains Back and Close navigation.','Every tab activates immediately by pointer and keyboard and displays only its own content.','Assignment tasks cannot be created without a title and due date.','Research milestones cannot be created without a title, owner and due date once ownership is implemented.','Deadline health is calculated consistently and displayed without overlap.','Next Action always points to a valid outstanding action.','Unrelated evidence never appears in assignment or research support panels.','Role restrictions are enforced by the API and reflected in the interface.','Notifications stop appearing urgent after the required action is completed.','TypeScript and production builds pass, and automated workflow tests cover critical state transitions.'])

h('14. Recommended Implementation Roadmap')
table(['Phase','Focus','Principal deliverables'],[
('1','Workspace foundations','Full-screen shells, single-panel tabs, identity headers and responsive navigation'),
('2','Assignment workflow','Task deadlines, ownership, progress, outputs, review and completion gates'),
('3','Research planning','Plan fields, approval, milestone ownership and research readiness'),
('4','Evidence governance','Source quality, provenance, links, citations and dataset controls'),
('5','Document collaboration','Section ownership, versioning, review, comments and final output controls'),
('6','Notifications and Felix','State-aware alerts, action resolution and permission-aware assistance'),
('7','Assurance','Automated tests, accessibility, audit validation, performance and release readiness')],[900,2500,5960])

h('15. Decision Summary')
p('App2 should use a common workspace design language while preserving the operational distinction between assignments and research. Assignment management is accountable delivery. Research management is controlled inquiry and evidence production. Linking them provides end-to-end traceability from an institutional request, through research, to an approved deliverable without weakening either governance model.')
callout('PRODUCT DIRECTION','Build shared navigation, identity, collaboration and audit patterns once; specialize workflow, evidence and completion controls for the purpose of each workspace.')

doc.core_properties.title='App2 Assignment and Research Workspaces Functional Specification'
doc.core_properties.subject='Functional specification and operating model'
doc.core_properties.author='Public Service Commission - App2 Project'
doc.save(OUT)
print(OUT)
